import ollama
import os
import fitz  # PyMuPDF para PDF
from docx import Document

# Modelo y carpeta
modelo = "llama3.2"
carpeta = "./Papers"

# Función para extraer texto de PDF
def extraer_texto_pdf(ruta):
    texto = ""
    with fitz.open(ruta) as doc:
        for pagina in doc:
            texto += pagina.get_text()
    return texto

# Función para enviar prompt a Ollama
def procesar_paper(texto):
    prompt = f"""
A continuación tienes el texto de un artículo científico sobre calidad del agua. Necesito que analices y respondas lo siguiente, Dame resultdados que no sean superficiales:

1. ¿Cuál es un resumen general del artículo?
2. ¿Qué problema trata de solucionar el paper?
3. ¿Cómo puede justificar este árticulo la realización de un sistema multi-sensor que mide pH, Turbiedad, Turbidez, Conductividad y ORP a tiempo real de manera remota usando IoT?
4. ¿Cómo se aborda la medición de calidad del agua en el estudio?
5. ¿Qué sensores se utilizan para las mediciones? Explica por qué se eligieron esos sensores en particular.
6. ¿Qué técnicas se emplean para la fusión de datos en el artículo? Especifica los enfoques usados para combinar los datos de diferentes fuentes.
7. ¿Qué tan eficaz fue el método o los modelos utilizados? Proporciona información sobre la validación o los resultados que midan la efectividad del enfoque.


Texto del artículo:
{texto[:4000]}  # Límite para no exceder tokens
"""

    response = ollama.chat(
        model=modelo,
        messages=[{"role": "user", "content": prompt}]
    )
    return response['message']['content']

# Crear documento Word
doc = Document()
doc.add_heading("Resumen de Papers sobre Calidad del Agua", 0)

# Procesar cada archivo
for archivo in os.listdir(carpeta):
    if archivo.endswith(".pdf"):
        ruta = os.path.join(carpeta, archivo)
        texto = extraer_texto_pdf(ruta)
        resumen = procesar_paper(texto)

        doc.add_heading(f"Archivo: {archivo}", level=1)
        doc.add_paragraph(resumen)
        doc.add_page_break()

# Guardar el archivo
doc.save("Resumen_Papers_CalidadAgua.docx")
print("✅ Resumen generado en 'Resumen_Papers_CalidadAgua.docx'")
